import struct
import zlib

from docx import Document

from agents.storyboard.docx_export import export_docx
from agents.storyboard.schemas import Shot, StoryboardOutput


def _minimal_png_1x1() -> bytes:
    """python-docx의 PNG 파서가 실제로 통과하는, 유효한 1x1 그레이스케일 PNG를 생성한다."""

    def chunk(chunk_type: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + chunk_type
            + data
            + struct.pack(">I", zlib.crc32(chunk_type + data))
        )

    signature = b"\x89PNG\r\n\x1a\n"
    ihdr = chunk(b"IHDR", struct.pack(">IIBBBBB", 1, 1, 8, 0, 0, 0, 0))
    raw_scanline = b"\x00\x00"  # filter byte(0) + 1 grayscale pixel
    idat = chunk(b"IDAT", zlib.compress(raw_scanline))
    iend = chunk(b"IEND", b"")
    return signature + ihdr + idat + iend

STORYBOARD = StoryboardOutput(
    shots=[
        Shot(
            size="ws",
            angle="eye-level",
            movement="static",
            duration=4,
            description="낡은 극장 무대 위, 지우가 소품 상자를 뒤진다",
            dialogue="",
            audio="새소리",
            sceneSlug="S1",
        ),
        Shot(
            size="ms",
            angle="low",
            movement="tilt-up",
            duration=3,
            description="지우가 고개를 든다",
            dialogue="아직 여기 있었네.",
            audio="",
            sceneSlug="S1",
        ),
    ]
)


def test_export_docx_creates_file_with_header_and_one_row_per_shot(tmp_path) -> None:
    path = tmp_path / "storyboard.docx"

    export_docx(STORYBOARD, path)

    assert path.exists()
    document = Document(str(path))
    table = document.tables[0]
    assert len(table.rows) == 1 + len(STORYBOARD.shots)
    header_texts = [cell.text for cell in table.rows[0].cells]
    assert header_texts == ["Cut", "Video", "Content", "Audio", "Time"]


def test_export_docx_converts_shot_notation_to_team_convention(tmp_path) -> None:
    path = tmp_path / "storyboard.docx"

    export_docx(STORYBOARD, path)

    document = Document(str(path))
    table = document.tables[0]
    second_shot_content = table.rows[2].cells[2].text
    assert "MS, Low Angle, Tilt Up" in second_shot_content
    assert "지우가 고개를 든다" in second_shot_content


def test_export_docx_puts_duration_in_time_column(tmp_path) -> None:
    path = tmp_path / "storyboard.docx"

    export_docx(STORYBOARD, path)

    document = Document(str(path))
    table = document.tables[0]
    assert table.rows[1].cells[4].text == "4"
    assert table.rows[2].cells[4].text == "3"


def test_export_docx_embeds_contact_sheet_image_when_available(tmp_path) -> None:
    path = tmp_path / "storyboard.docx"
    images_dir = tmp_path / "images"
    images_dir.mkdir()
    (images_dir / "cut1.png").write_bytes(_minimal_png_1x1())

    export_docx(STORYBOARD, path, images_dir=images_dir)

    document = Document(str(path))
    assert len(document.inline_shapes) == 1
