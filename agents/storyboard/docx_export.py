"""팀 표준 5컬럼 워드 출력 (기획서 4.8절)."""
from pathlib import Path

from docx import Document
from docx.shared import Inches, Twips

from agents.storyboard.schemas import Shot, StoryboardOutput

COLUMN_HEADERS = ["Cut", "Video", "Content", "Audio", "Time"]
COLUMN_WIDTHS_DXA = [690, 2925, 3030, 1530, 840]
IMAGE_WIDTH_INCHES = 1.89

ANGLE_LABELS = {
    "overhead": "Top Angle",
    "low": "Low Angle",
    "high": "High Angle",
    "dutch": "Dutch Angle",
    "pov": "POV",
    "ots": "OSS",
}

MOVEMENT_LABELS = {
    "pan": "Pan",
    "pan-left": "Pan Left",
    "pan-right": "Pan Right",
    "tilt": "Tilt",
    "tilt-up": "Tilt Up",
    "tilt-down": "Tilt Down",
    "dolly-in": "Dolly In",
    "dolly-out": "Dolly Out",
    "zoom-in": "Zoom In",
    "zoom-out": "Zoom Out",
    "tracking": "Tracking",
    "handheld": "Handheld",
    "crane": "Crane",
    "zoom": "Zoom",
    "arc": "Arc",
}


def shot_notation(shot: Shot) -> str:
    """팀 고유 표기 관례로 변환한다(기획서 4.8-3절): eye-level/static은 생략."""
    parts = [shot.size.upper()]
    angle_label = ANGLE_LABELS.get(shot.angle.lower())
    if angle_label:
        parts.append(angle_label)
    movement_label = MOVEMENT_LABELS.get(shot.movement.lower())
    if movement_label:
        parts.append(movement_label)
    return ", ".join(parts)


def _set_column_widths(table) -> None:
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, COLUMN_WIDTHS_DXA):
            cell.width = Twips(width)


def export_docx(
    storyboard: StoryboardOutput, path: Path, images_dir: Path | None = None
) -> None:
    document = Document()
    for section in document.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    table = document.add_table(rows=1, cols=len(COLUMN_HEADERS))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for cell, header in zip(header_cells, COLUMN_HEADERS):
        cell.paragraphs[0].add_run(header).bold = True

    for i, shot in enumerate(storyboard.shots, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)

        image_path = (images_dir / f"cut{i}.png") if images_dir else None
        if image_path and image_path.exists():
            run = row_cells[1].paragraphs[0].add_run()
            run.add_picture(str(image_path), width=Inches(IMAGE_WIDTH_INCHES))

        content_cell = row_cells[2]
        content_cell.paragraphs[0].add_run(shot_notation(shot))
        content_cell.add_paragraph(shot.description)

        audio_parts = [p for p in (shot.dialogue, shot.audio) if p]
        row_cells[3].text = " / ".join(audio_parts)

        row_cells[4].text = str(shot.duration)

    _set_column_widths(table)
    document.save(str(path))
